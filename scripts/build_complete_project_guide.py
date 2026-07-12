"""Build the comprehensive technical guide from its Markdown source."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_report import MUTED, NAVY, BLUE, add_page_field, configure_styles, render_markdown, set_run_font


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "complete_project_guide.md"
OUTPUT = ROOT / "deliverables" / "advanced_rag_llm_complete_technical_guide.docx"


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("Advanced RAG and LLM APIs - Complete Technical Guide | Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(paragraph)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("COMPLETE TECHNICAL GUIDE")
    set_run_font(run, size=12, color=BLUE, bold=True)
    paragraph.paragraph_format.space_after = Pt(18)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Development of Advanced AI Applications using RAG and LLM APIs")
    set_run_font(run, size=26, color=NAVY, bold=True)
    paragraph.paragraph_format.space_after = Pt(14)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "Concepts, architecture, implementation, evaluation, operations, security, and glossary"
    )
    set_run_font(run, size=14, color=MUTED)
    paragraph.paragraph_format.space_after = Pt(56)

    for label, value in (
        ("Audience", "Master's students, supervisors, examiners, and developers"),
        ("System", "FastAPI, SQLite, hybrid retrieval, and pluggable LLM providers"),
        ("Version", "Project documentation generated July 2026"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lead = paragraph.add_run(f"{label}: ")
        set_run_font(lead, size=11, color=NAVY, bold=True)
        detail = paragraph.add_run(value)
        set_run_font(detail, size=11, color=MUTED)

    doc.add_page_break()

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Reading guide")
    for text in (
        "Sections 1 through 5 introduce the problem and foundational concepts.",
        "Sections 6 through 13 describe the architecture, code modules, and operating workflow.",
        "Sections 14 through 18 explain evaluation, verification, security, limitations, and the internship plan.",
        "Section 19 is a reference glossary of important terms used throughout the project.",
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(text), size=11)
    doc.add_page_break()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
    add_footer(section)
    add_cover(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
