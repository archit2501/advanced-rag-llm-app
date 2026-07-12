"""Build the final internship report from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "final_report.md"
OUTPUT = ROOT / "deliverables" / "advanced_rag_llm_internship_report.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 99, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, *, size: float, color: RGBColor = None, bold: bool = False, italic: bool = False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: List[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    table_properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("Advanced RAG and LLM APIs Internship Report | Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(paragraph)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = doc.styles.add_style("Report Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(9)
    code.paragraph_format.space_after = Pt(4)


def add_cover(doc: Document) -> None:
    for _ in range(7):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("SUMMER INTERNSHIP REPORT")
    set_run_font(run, size=12, color=BLUE, bold=True)
    paragraph.paragraph_format.space_after = Pt(20)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Development of Advanced AI Applications using RAG and LLM APIs")
    set_run_font(run, size=27, color=NAVY, bold=True)
    paragraph.paragraph_format.space_after = Pt(14)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("A grounded knowledge-assistant prototype with retrieval evaluation and LLM API integration")
    set_run_font(run, size=14, color=MUTED)
    paragraph.paragraph_format.space_after = Pt(70)

    for label, value in (
        ("Programme", "Master's Summer Internship"),
        ("Prepared", "July 2026"),
        ("System", "FastAPI, SQLite, hybrid retrieval, and pluggable LLM providers"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lead = paragraph.add_run(f"{label}: ")
        set_run_font(lead, size=11, color=NAVY, bold=True)
        detail = paragraph.add_run(value)
        set_run_font(detail, size=11, color=MUTED)
    doc.add_page_break()


def parse_inline(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(size - 0.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def split_table_row(line: str) -> List[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: List[List[str]]) -> None:
    if len(rows) < 2:
        return
    data = [rows[0]] + rows[2:]
    column_count = len(data[0])
    widths = [CONTENT_WIDTH_DXA // column_count] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(data):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            paragraph = cells[column_index].paragraphs[0]
            parse_inline(paragraph, value, size=9.5)
            if row_index == 0:
                set_cell_shading(cells[column_index], LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                set_cell_shading(cells[column_index], LIGHT_GRAY)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_markdown(doc: Document, lines: Iterable[str]) -> None:
    lines = list(lines)
    index = 0
    code_mode = False
    table_rows: List[List[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            flush_table()
            code_mode = not code_mode
            index += 1
            continue
        if code_mode:
            paragraph = doc.add_paragraph(style="Report Code")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            index += 1
            continue
        if line.startswith("|"):
            table_rows.append(split_table_row(line))
            index += 1
            continue
        flush_table()
        if not line:
            index += 1
            continue
        if line == "<!-- pagebreak -->":
            doc.add_page_break()
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("### "):
            if line.startswith("### 7.2 Retrieval comparison"):
                doc.add_page_break()
            paragraph = doc.add_paragraph(style="Heading 3")
            parse_inline(paragraph, line[4:], size=12)
        elif line.startswith("## "):
            if line.startswith("## 10. Six-Week Internship Timeline"):
                doc.add_page_break()
            paragraph = doc.add_paragraph(style="Heading 1")
            parse_inline(paragraph, line[3:], size=16)
        elif re.match(r"\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            parse_inline(paragraph, re.sub(r"^\d+\. ", "", line), size=11)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            parse_inline(paragraph, line[2:], size=11)
        else:
            paragraph = doc.add_paragraph()
            parse_inline(paragraph, line, size=11)
        index += 1
    flush_table()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_footer(section)
    add_cover(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
